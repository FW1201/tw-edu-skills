#!/usr/bin/env python3
"""差異化教學設計文件生成腳本"""
import argparse
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_BLUE  = RGBColor(0x1A, 0x52, 0x76)
C_MID   = RGBColor(0x24, 0x71, 0xA3)
C_LIGHT = RGBColor(0xEB, 0xF5, 0xFB)
C_GREEN = RGBColor(0x1E, 0x84, 0x49)
C_GOLD  = RGBColor(0xD4, 0xAC, 0x0D)
C_ORANGE= RGBColor(0xCA, 0x6F, 0x1E)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK  = RGBColor(0x1C, 0x2A, 0x35)

def rgb_hex(c): return '{:02X}{:02X}{:02X}'.format(c[0], c[1], c[2])

def set_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    tcPr.append(shd)

def set_border(cell, color='2471A3'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, size=11, color=C_DARK, center=False):
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    r.font.name = '標楷體'; r.font.color.rgb = color
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rPr = r._r.get_or_add_rPr()
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), '標楷體')
    rPr.insert(0, rF)

def hdr_cell(cell, text, bg=C_MID):
    set_bg(cell, bg); set_border(cell)
    cell_text(cell, text, bold=True, color=C_WHITE, center=True)

def data_cell(cell, text, row=0, center=False):
    bg = C_LIGHT if row % 2 == 0 else RGBColor(0xF8, 0xF9, 0xFA)
    set_bg(cell, bg); set_border(cell)
    cell_text(cell, text, center=center)

def section_title(doc, text, color=C_BLUE):
    p = doc.add_paragraph()
    r = p.add_run(f'■ {text}')
    r.bold = True; r.font.size = Pt(13)
    r.font.name = '標楷體'; r.font.color.rgb = color
    rPr = r._r.get_or_add_rPr()
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), '標楷體')
    rPr.insert(0, rF)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_level_table(doc, subject, title, grade):
    """三層次學習任務對照表"""
    section_title(doc, '表格一：三層次學習任務設計')
    tbl = doc.add_table(rows=4, cols=5)
    tbl.style = 'Table Grid'
    widths = [Cm(2.5), Cm(4), Cm(5), Cm(4), Cm(3.5)]
    for j, w in enumerate(widths):
        tbl.columns[j].width = w

    hdrs = ['學習層次', '學習目標', '學習活動', '支援策略', '評量調整']
    level_configs = [
        ('基礎層\n（Essential）', C_ORANGE,
         f'·能說出{title}的主要內容\n·完成關鍵概念填空',
         f'·提供文章摘要單\n·有鷹架的問題引導\n·允許查閱字典/筆記',
         '·降低題目難度\n·延長作答時間\n·允許口頭作答'),
        ('標準層\n（Core）', C_MID,
         f'·能分析{title}的結構與主旨\n·完成標準學習任務',
         '·教師提示重點\n·同儕學習配對\n·適度提問引導',
         '·依課綱標準評量\n·學習單 + 口頭'),
        ('進階層\n（Advanced）', C_GREEN,
         f'·能評鑑{title}的寫作技巧\n·完成開放性延伸任務',
         '·開放性探究任務\n·自主選擇表達方式\n·連結跨域資料',
         '·高層次思維題\n·創意作品評量'),
    ]
    for j, h in enumerate(hdrs):
        hdr_cell(tbl.rows[0].cells[j], h, bg=C_BLUE)

    for i, (level, color, activity, support, assess) in enumerate(level_configs, 1):
        row = tbl.rows[i]
        hdr_cell(row.cells[0], level, bg=color)
        for j, text in enumerate([
            f'依{grade}課綱調整\n（布魯姆 C1-C2）' if i==1 else
            f'符合{grade}課綱標準\n（布魯姆 C2-C4）' if i==2 else
            f'超越課綱深化\n（布魯姆 C4-C6）',
            activity, support, assess
        ], 1):
            data_cell(row.cells[j], text, row=i)
    doc.add_paragraph()

def add_udl_table(doc):
    """UDL 三原則應用表"""
    section_title(doc, '表格二：UDL 通用設計學習應用')
    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = 'Table Grid'
    for j, w in enumerate([Cm(5), Cm(5.5), Cm(8.5)]):
        tbl.columns[j].width = w

    hdrs = ['UDL 原則', '設計重點', '具體做法範例']
    for j, h in enumerate(hdrs):
        hdr_cell(tbl.rows[0].cells[j], h, bg=C_BLUE)

    rows_data = [
        ('一、提供多元\n表徵方式\n（What of Learning）',
         '以不同形式呈現學習內容，確保每位學生都能接收訊息',
         '·課文文字 + 朗讀音檔\n·提供圖表輔助理解\n·關鍵詞彙視覺化\n·提供注音版輔助讀物'),
        ('二、提供多元\n行動與表達\n（How of Learning）',
         '讓學生用不同方式展現學習成果，不侷限單一形式',
         '·書寫作答 or 口頭報告\n·個人作業 or 小組合作\n·選擇自己擅長的呈現方式\n·允許使用輔助工具'),
        ('三、提供多元\n參與方式\n（Why of Learning）',
         '透過不同策略維持學習動機與投入感',
         '·連結個人生活經驗\n·提供選擇學習主題的自主權\n·設計小組合作任務\n·建立安全正向的學習氛圍'),
    ]
    for i, (principle, focus, example) in enumerate(rows_data, 1):
        data_cell(tbl.rows[i].cells[0], principle, row=i, center=True)
        data_cell(tbl.rows[i].cells[1], focus, row=i)
        data_cell(tbl.rows[i].cells[2], example, row=i)
    doc.add_paragraph()

def add_special_needs_table(doc, needs_text):
    """個別學生調整建議"""
    section_title(doc, '表格三：特殊需求學生調整建議')
    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = 'Table Grid'
    for j, w in enumerate([Cm(4), Cm(4), Cm(5), Cm(6)]):
        tbl.columns[j].width = w

    hdrs = ['學生類型', '主要需求特性', '課堂調整建議', '評量調整建議']
    for j, h in enumerate(hdrs):
        hdr_cell(tbl.rows[0].cells[j], h, bg=C_BLUE)

    special_data = [
        ('閱讀障礙學生', '文字解碼困難、閱讀速度慢',
         '·提供大字版或注音版\n·允許使用文字轉語音\n·給予額外時間',
         '·口頭問答取代筆試\n·減少書寫量\n·允許使用輔具'),
        ('資優/高能力學生', '學習速度快、需要更多挑戰',
         '·提供延伸思考問題\n·允許獨立探究\n·指定進階讀本',
         '·設計開放性評量\n·自選創作形式\n·加分制度'),
        ('新住民/華語學習者', '中文詞彙量不足',
         '·提供關鍵詞彙翻譯對照\n·允許以母語輔助思考\n·提供圖文並茂的輔助材料',
         '·降低語言難度要求\n·允許雙語作答\n·著重內容理解'),
    ]
    for i, row_data in enumerate(special_data, 1):
        for j, text in enumerate(row_data):
            data_cell(tbl.rows[i].cells[j], text, row=i)
    doc.add_paragraph()

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--subject',  default='國語文')
    parser.add_argument('--title',    default='課文')
    parser.add_argument('--grade',    default='國中八年級')
    parser.add_argument('--needs',    default='含閱讀障礙學生2名、資優生1名')
    parser.add_argument('--output',   default='差異化教學設計.docx')
    args = parser.parse_args()

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # 封面
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('差異化教學設計方案')
    r.bold = True; r.font.size = Pt(24)
    r.font.name = '標楷體'; r.font.color.rgb = C_BLUE
    rPr = r._r.get_or_add_rPr()
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), '標楷體')
    rPr.insert(0, rF)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'《{args.title}》｜{args.subject}｜{args.grade}')
    r2.font.size = Pt(16); r2.font.name = '標楷體'
    r2.font.color.rgb = C_MID
    rPr2 = r2._r.get_or_add_rPr()
    rF2 = OxmlElement('w:rFonts')
    rF2.set(qn('w:eastAsia'), '標楷體')
    rPr2.insert(0, rF2)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f'班級特性：{args.needs}')
    r3.font.size = Pt(12); r3.font.name = '標楷體'
    r3.font.color.rgb = C_DARK
    doc.add_page_break()

    add_level_table(doc, args.subject, args.title, args.grade)
    add_udl_table(doc)
    add_special_needs_table(doc, args.needs)

    doc.save(args.output)
    print(f'✓ 差異化教學設計已儲存：{args.output}')

if __name__ == '__main__':
    main()
