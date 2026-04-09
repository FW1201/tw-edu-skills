#!/usr/bin/env python3
"""
臺灣 108 課綱素養導向教案生成腳本
用法: python generate_lesson_plan.py --subject 國語文 --title 背影 --grade 國中八年級 ...
"""

import argparse
import sys
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 色彩常數（臺灣教育主題色）──────────────────────────
BLUE_DARK   = RGBColor(0x1A, 0x52, 0x76)   # #1A5276 深藍（主色）
BLUE_MID    = RGBColor(0x24, 0x71, 0xA3)   # #2471A3 中藍（表頭）
BLUE_LIGHT  = RGBColor(0xEB, 0xF5, 0xFB)   # #EBF5FB 淺藍（奇數列）
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LIGHT  = RGBColor(0xF8, 0xF9, 0xFA)   # 偶數列
GOLD        = RGBColor(0xD4, 0xAC, 0x0D)   # 強調色
TEXT_DARK   = RGBColor(0x1C, 0x2A, 0x35)   # 深色文字

# ── 工具函式 ──────────────────────────────────────────

def set_cell_bg(cell, color: RGBColor):
    """設定儲存格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=True, bottom=True, left=True, right=True,
                    color='2471A3', size='4'):
    """設定儲存格框線"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, active in [('top', top), ('bottom', bottom),
                          ('left', left), ('right', right)]:
        if active:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), size)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_cell_text(cell, text, bold=False, font_size=11,
                  color: RGBColor = TEXT_DARK, center=False):
    """向儲存格加入文字"""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.name = '標楷體'
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 設定東亞字型
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '標楷體')
    rFonts.set(qn('w:ascii'), 'Arial')
    rPr.insert(0, rFonts)
    return para

def make_header_cell(cell, text):
    """建立深藍底白字的表頭儲存格"""
    set_cell_bg(cell, BLUE_MID)
    set_cell_border(cell, color='1A5276', size='6')
    add_cell_text(cell, text, bold=True, font_size=11,
                  color=WHITE, center=True)

def make_data_cell(cell, text, row_idx=0, center=False):
    """建立資料儲存格（奇偶列交替底色）"""
    bg = BLUE_LIGHT if row_idx % 2 == 0 else GRAY_LIGHT
    set_cell_bg(cell, bg)
    set_cell_border(cell, color='2471A3', size='4')
    add_cell_text(cell, text, font_size=11,
                  color=TEXT_DARK, center=center)

def add_section_title(doc, title: str, level: int = 1):
    """加入帶底線的章節標題"""
    para = doc.add_paragraph()
    para.clear()
    run = para.add_run(f'▌ {title}')
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.name = '標楷體'
    run.font.color.rgb = BLUE_DARK
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '2471A3')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    # 東亞字型
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '標楷體')
    rPr.insert(0, rFonts)

def add_cover_page(doc, subject, title, grade, publisher,
                   periods, teacher, school=''):
    """加入美觀封面頁"""
    # 標題段落
    doc.add_paragraph()
    doc.add_paragraph()

    # 學校名稱
    if school:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(school)
        run.font.size = Pt(16)
        run.font.name = '標楷體'
        run.font.color.rgb = BLUE_DARK
        _set_eastasia(run)

    doc.add_paragraph()

    # 主標題（教案名稱）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('素養導向教案')
    run.font.size = Pt(28)
    run.bold = True
    run.font.name = '標楷體'
    run.font.color.rgb = BLUE_DARK
    _set_eastasia(run)

    # 副標題（課文/主題）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'《{title}》')
    run.font.size = Pt(22)
    run.bold = True
    run.font.name = '標楷體'
    run.font.color.rgb = BLUE_MID
    _set_eastasia(run)

    doc.add_paragraph()
    doc.add_paragraph()

    # 基本資料表格
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    info = [
        ['領域/科目', subject, '適用年級', grade],
        ['教材版本', publisher, '教學節數', f'{periods} 節（每節 45 分鐘）'],
        ['設計者', teacher, '設計日期', str(date.today())],
        ['學校', school if school else '　', '班級特性', '普通班'],
    ]
    col_widths = [Cm(3), Cm(6), Cm(3), Cm(6)]
    for i, row_data in enumerate(info):
        row = table.rows[i]
        for j, (text, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[j]
            cell.width = w
            if j % 2 == 0:
                make_header_cell(cell, text)
            else:
                make_data_cell(cell, text, row_idx=i)

    doc.add_paragraph()
    doc.add_page_break()

def _set_eastasia(run):
    """設定 run 的東亞字型"""
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '標楷體')
    rFonts.set(qn('w:ascii'), 'Arial')
    rPr.insert(0, rFonts)

# ── 各表格生成函式 ──────────────────────────────────

def add_table1_basic(doc, data: dict):
    """表格一：基本資料"""
    add_section_title(doc, '表格一：基本資料')
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    col_widths = [Cm(3), Cm(5.5), Cm(3), Cm(5.5)]
    rows_data = [
        ['領域/科目', data.get('subject', ''), '學習階段', data.get('stage', '')],
        ['單元/課文名稱', data.get('title', ''), '設計者', data.get('teacher', '')],
        ['教材版本', data.get('publisher', ''), '教學節數',
         f"{data.get('periods', 3)} 節（每節 45 分鐘）"],
    ]
    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths[j]
            if j % 2 == 0:
                make_header_cell(cell, text)
            else:
                make_data_cell(cell, text, row_idx=i)
    doc.add_paragraph()

def add_table2_curriculum(doc, competencies: list, performance: list, content: list):
    """表格二：設計依據（108課綱對應）"""
    add_section_title(doc, '表格二：設計依據（108 課綱對應）')
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(13.5)

    # 表頭
    make_header_cell(table.rows[0].cells[0], '類別')
    make_header_cell(table.rows[0].cells[1], '代碼與說明')

    rows_data = [
        ('核心素養', '\n'.join(competencies)),
        ('學習表現', '\n'.join(performance)),
        ('學習內容', '\n'.join(content)),
    ]
    for i, (label, text) in enumerate(rows_data, 1):
        make_data_cell(table.rows[i].cells[0], label, row_idx=i, center=True)
        make_data_cell(table.rows[i].cells[1], text, row_idx=i)
    doc.add_paragraph()

def add_table3_objectives(doc, obj_cog: str, obj_aff: str, obj_ski: str):
    """表格三：學習目標"""
    add_section_title(doc, '表格三：學習目標')
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(13.5)
    make_header_cell(table.rows[0].cells[0], '面向')
    make_header_cell(table.rows[0].cells[1], '學習目標')
    rows_data = [('認知面', obj_cog), ('情意面', obj_aff), ('技能面', obj_ski)]
    for i, (label, text) in enumerate(rows_data, 1):
        make_data_cell(table.rows[i].cells[0], label, row_idx=i, center=True)
        make_data_cell(table.rows[i].cells[1], text, row_idx=i)
    doc.add_paragraph()

def add_table4_analysis(doc, text_type: str, theme: str, structure: str,
                        technique: str, culture: str):
    """表格四：課文/主題分析"""
    add_section_title(doc, '表格四：課文分析')
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(13.5)
    make_header_cell(table.rows[0].cells[0], '分析項目')
    make_header_cell(table.rows[0].cells[1], '說明')
    rows_data = [
        ('文體', text_type),
        ('主題意旨', theme),
        ('篇章結構', structure),
        ('表達技巧', technique),
        ('文化意涵', culture),
    ]
    for i, (label, text) in enumerate(rows_data, 1):
        make_data_cell(table.rows[i].cells[0], label, row_idx=i, center=True)
        make_data_cell(table.rows[i].cells[1], text, row_idx=i)
    doc.add_paragraph()

def add_table5_vocab(doc, vocab_list: list):
    """表格五：生難字詞（國語文用）"""
    if not vocab_list:
        return
    add_section_title(doc, '表格五：生難字詞')
    table = doc.add_table(rows=len(vocab_list) + 1, cols=5)
    table.style = 'Table Grid'
    col_widths = [Cm(2), Cm(2.5), Cm(1.5), Cm(5), Cm(6)]
    headers = ['字詞', '注音', '詞性', '解釋', '例句']
    for j, (h, w) in enumerate(zip(headers, col_widths)):
        table.columns[j].width = w
        make_header_cell(table.rows[0].cells[j], h)
    for i, item in enumerate(vocab_list, 1):
        row = table.rows[i]
        for j, text in enumerate(item):
            make_data_cell(row.cells[j], str(text), row_idx=i, center=(j < 3))
    doc.add_paragraph()

def add_table6_activities(doc, periods_data: list):
    """表格六：教學活動設計（多節課）"""
    for period_num, period in enumerate(periods_data, 1):
        add_section_title(doc,
            f'表格六：教學活動設計（第 {period_num} 節）', level=1)

        # 節次說明
        p = doc.add_paragraph()
        run = p.add_run(f'⏱ 第 {period_num} 節・共 45 分鐘')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = BLUE_MID
        run.font.name = '標楷體'
        _set_eastasia(run)

        table = doc.add_table(rows=len(period['rows']) + 1, cols=4)
        table.style = 'Table Grid'
        col_widths = [Cm(2.5), Cm(9.5), Cm(1.5), Cm(3.5)]
        headers = ['教學流程', '教學活動詳細步驟', '時間', '評量方式']
        for j, (h, w) in enumerate(zip(headers, col_widths)):
            table.columns[j].width = w
            make_header_cell(table.rows[0].cells[j], h)
        for i, row_data in enumerate(period['rows'], 1):
            row = table.rows[i]
            make_data_cell(row.cells[0], row_data.get('flow', ''), row_idx=i, center=True)
            make_data_cell(row.cells[1], row_data.get('activity', ''), row_idx=i)
            make_data_cell(row.cells[2], row_data.get('time', ''), row_idx=i, center=True)
            make_data_cell(row.cells[3], row_data.get('assess', ''), row_idx=i)
        doc.add_paragraph()

def add_table7_literacy(doc, rows: list):
    """表格七：素養導向教學重點"""
    add_section_title(doc, '表格七：素養導向教學重點說明')
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.style = 'Table Grid'
    col_widths = [Cm(3.5), Cm(7), Cm(6.5)]
    headers = ['素養面向', '具體教學做法', '課文/主題連結說明']
    for j, (h, w) in enumerate(zip(headers, col_widths)):
        table.columns[j].width = w
        make_header_cell(table.rows[0].cells[j], h)
    for i, row_data in enumerate(rows, 1):
        make_data_cell(table.rows[i].cells[0], row_data.get('dimension', ''), row_idx=i, center=True)
        make_data_cell(table.rows[i].cells[1], row_data.get('method', ''), row_idx=i)
        make_data_cell(table.rows[i].cells[2], row_data.get('link', ''), row_idx=i)
    doc.add_paragraph()

def add_table8_assessment(doc, rows: list):
    """表格八：評量規劃"""
    add_section_title(doc, '表格八：評量規劃')
    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = 'Table Grid'
    col_widths = [Cm(4), Cm(4), Cm(3), Cm(6)]
    headers = ['學習目標', '評量方式', '評量時機', '評量規準']
    for j, (h, w) in enumerate(zip(headers, col_widths)):
        table.columns[j].width = w
        make_header_cell(table.rows[0].cells[j], h)
    for i, row_data in enumerate(rows, 1):
        for j, key in enumerate(['objective', 'method', 'timing', 'criteria']):
            make_data_cell(table.rows[i].cells[j], row_data.get(key, ''), row_idx=i)
    doc.add_paragraph()

def add_table9_extension(doc, rows: list):
    """表格九：延伸學習"""
    add_section_title(doc, '表格九：延伸學習')
    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = 'Table Grid'
    col_widths = [Cm(3.5), Cm(5), Cm(3.5), Cm(5)]
    headers = ['延伸主題', '建議學習資源', '跨域連結', '對應重大議題']
    for j, (h, w) in enumerate(zip(headers, col_widths)):
        table.columns[j].width = w
        make_header_cell(table.rows[0].cells[j], h)
    for i, row_data in enumerate(rows, 1):
        for j, key in enumerate(['theme', 'resource', 'cross', 'issue']):
            make_data_cell(table.rows[i].cells[j], row_data.get(key, ''), row_idx=i)
    doc.add_paragraph()


# ── 主程式 ──────────────────────────────────────────

def build_sample_data(args) -> dict:
    """建立範例資料（當 Claude 實際調用時應替換為真實內容）"""
    grade_stage_map = {
        '國小一年': '第一學習階段', '國小二年': '第一學習階段',
        '國小三年': '第二學習階段', '國小四年': '第二學習階段',
        '國小五年': '第三學習階段', '國小六年': '第三學習階段',
        '國中七年': '第四學習階段', '國中八年': '第四學習階段',
        '國中九年': '第四學習階段',
        '高一': '第五學習階段', '高二': '第五學習階段', '高三': '第五學習階段',
    }
    stage = next((v for k, v in grade_stage_map.items()
                  if k in args.grade), '第四學習階段')

    return {
        # 基本資料
        'subject': args.subject,
        'title': args.title,
        'grade': args.grade,
        'publisher': args.publisher,
        'periods': args.periods,
        'teacher': args.teacher,
        'school': args.school,
        'stage': stage,

        # 課綱對應（範例，Claude 應替換為真實對應）
        'competencies': [
            '語-J-B1：運用國語文表情達意，增進閱讀理解，提升欣賞及評析能力。',
            '語-J-A2：透過欣賞文學作品，培養思辨能力，建構正向價值觀。',
            '語-J-C1：透過文學作品認識倫理課題，以適切態度與人互動。',
        ],
        'performance': [
            '5-Ⅳ-2：能理解並分析文本中的語句、段落、篇章結構及寫作手法。',
            '5-Ⅳ-5：能閱讀不同時代的文學作品，感受其文化意涵。',
            '6-Ⅳ-2：能運用修辭策略，增強表達效果。',
        ],
        'content': [
            'Ab-Ⅳ-1：詞語積累，同義詞、反義詞的辨析與運用。',
            'Ac-Ⅳ-3：記敘文的時間、空間、人物與情節安排。',
            'Ca-Ⅳ-1：修辭技巧的運用（動作描寫、細節描寫）。',
            'Da-Ⅳ-1：重要現代文學作品的時代背景與文化內涵。',
        ],

        # 學習目標
        'obj_cog': f'1. 學生能正確說出《{args.title}》中 10 個生難字詞的注音與字義。\n2. 學生能分析課文的篇章結構與記敘脈絡。\n3. 學生能辨識並說明課文中的修辭手法與表達技巧。',
        'obj_aff': f'1. 學生能感受作者透過細節描寫傳達的情感。\n2. 學生能連結個人生命經驗，反思親情意涵。',
        'obj_ski': f'1. 學生能模仿課文的動作描寫方式，寫出 100 字以上的記敘片段。\n2. 學生能在小組討論中清楚表達個人觀點。',

        # 課文分析
        'text_type': '現代散文・記敘文',
        'theme': '透過父親爬月台買橘子的背影，呈現深沉的父子之愛與生命省思。',
        'structure': '起（交代背景）→ 承（車站送行）→ 轉（背影特寫）→ 合（事後回憶）',
        'technique': '細節描寫（動作描寫）、白描手法、第一人稱敘事、情景交融',
        'culture': '五四時代父子關係的重新詮釋；現代散文「真情流露」的典範之作',

        # 字詞
        'vocab': [
            ['蹣跚', 'ㄆㄢˊ ㄕㄢ', '形容詞', '走路搖晃、步履不穩的樣子', '老人蹣跚地走下階梯。'],
            ['狼狽', 'ㄌㄤˊ ㄅㄟˋ', '形容詞', '形容困苦或窘迫的樣子', '遭遇挫折後，他顯得十分狼狽。'],
            ['頹唐', 'ㄊㄨㄟˊ ㄊㄤˊ', '形容詞', '意志消沉、精神萎靡的樣子', '連遭失敗，他神情頹唐。'],
            ['禍不單行', 'ㄏㄨㄛˋ ㄅㄨˋ ㄉㄢ ㄒㄧㄥˊ', '成語', '壞事接連發生，接踵而來', '那年家中禍不單行，父親又失業了。'],
            ['交卸', 'ㄐㄧㄠ ㄒㄧㄝˋ', '動詞', '把職務移交給接任者', '他辦完交卸手續才返鄉。'],
        ],

        # 教學活動（三節課）
        'periods_data': [
            {
                'rows': [
                    {
                        'flow': '準備活動\n（5分）',
                        'activity': '【引起動機】\n1. 教師投影一組「父親或長輩的背影」照片（車站月台、送別場景），靜默約 30 秒。\n2. 提問：「你有沒有記憶中某個家人特定的動作或姿態，讓你印象深刻？當時是什麼情境？」請 2-3 位學生分享。\n3. 教師收束引入：「今天的散文，作者正是因為一個平凡的背影，重新認識了父親的愛。讓我們一起去找那個背影。」',
                        'time': '5分',
                        'assess': '口頭觀察（情感連結與參與意願）',
                    },
                    {
                        'flow': '發展活動\n（35分）',
                        'activity': '【一、字詞教學】（15分）\n1. 教師板書 5 個生難字詞（蹣跚、狼狽、頹唐、禍不單行、交卸），逐一附注音，朗讀 2 遍。\n2. 提問：「『蹣跚』的部首是足字旁，想想看，這跟什麼動作有關？」引導學生從字形猜義。\n3. 學生完成字詞學習單（填空 + 選詞填空），同桌互核。\n\n【二、作者與背景導入】（10分）\n1. 投影朱自清照片，簡介作者（1898-1948）及寫作背景（1925 年）。\n2. 說明家庭背景：父子關係曾一度疏離。提問：「如果你和父母之間有過誤解，什麼事讓你重新靠近彼此？」\n\n【三、初讀課文】（10分）\n1. 教師範讀全文，學生跟讀，並用符號標記：□不懂、★最感動的句子。\n2. 學生默讀後，用 2 句話說出文章大意，指名回答，教師板書：背影／父親／送行／愛。\n3. 提問：「文章中有四次提到『背影』，你覺得每次出現的意義一樣嗎？」',
                        'time': '35分',
                        'assess': '字詞學習單（形成性）；口頭問答觀察',
                    },
                    {
                        'flow': '綜合活動\n（5分）',
                        'activity': '【小結與預告】\n1. 提問：「今天初讀，哪個場景最讓你有感觸？」1-2 位分享。\n2. 預告第二節：「下節我們深入分析父親爬月台那段的動作描寫。」\n3. 課後任務：找出課文中最令你感動的一句話，寫下原因 3-5 行，下節課前繳交。',
                        'time': '5分',
                        'assess': '課後任務說明',
                    },
                ],
            },
            {
                'rows': [
                    {
                        'flow': '準備活動\n（5分）',
                        'activity': '【複習連結】\n1. 抽查 3-4 名學生分享課後感動句子及理由。\n2. 教師板書上節關鍵詞，提問：「上節課我們找到了背影，今天我們要問：朱自清為什麼會流淚？」',
                        'time': '5分',
                        'assess': '課後任務收回',
                    },
                    {
                        'flow': '發展活動\n（35分）',
                        'activity': '【一、精讀核心場景：父親爬月台】（20分）\n1. 教師朗讀父親爬月台段落，學生圈出所有動詞：探、穿、爬、攀、縮、傾。\n2. 提問：「這六個動詞，每一個都在告訴我們什麼？父親的身體狀況怎麼樣？」小組討論 3 分鐘，代表分享。\n3. 提問：「朱自清為什麼在父親爬月台的時候笑了自己聰明？這句話背後有什麼情感？」\n4. 教師補充：白描手法的特點——不加修飾，用精準的動作細節讓讀者自己感受。\n\n【二、文章結構梳理】（10分）\n1. 教師引導學生填寫結構分析學習單：四次提到背影的位置與各自的作用。\n2. 全班核對，教師板書結構圖：前兩次（伏筆）→ 第三次（高潮）→ 第四次（回憶）\n\n【三、修辭與情感討論】（5分）\n提問：「如果把這段動作描寫換成直接說『父親很辛苦地去買了橘子』，效果有什麼不同？」',
                        'time': '35分',
                        'assess': '結構分析學習單（小組）；口頭討論觀察',
                    },
                    {
                        'flow': '綜合活動\n（5分）',
                        'activity': '【深層省思】\n1. 教師問：「你認為作者寫這篇文章的目的是什麼？僅僅是紀念父親，還是還有更深的意義？」\n2. 預告第三節：「下節課，我們要用朱自清的方法——細節描寫——來寫我們自己的親情故事。」',
                        'time': '5分',
                        'assess': '口頭觀察',
                    },
                ],
            },
            {
                'rows': [
                    {
                        'flow': '準備活動\n（5分）',
                        'activity': '【仿作暖身】\n1. 回顧白描手法：選一個精準動詞 > 寫一個細節動作的小練習（2分鐘快速練習）。\n2. 全班分享 2-3 個句子，教師給予即時回饋。',
                        'time': '5分',
                        'assess': '口頭觀察',
                    },
                    {
                        'flow': '發展活動\n（30分）',
                        'activity': '【一、仿作寫作】（20分）\n1. 任務說明：模仿課文的動作細節描寫，以「某位家人的某個動作或姿態」為題，寫 100-150 字的記敘片段。要求：至少使用 3 個精準動詞；不直接說出情感，讓動作說話。\n2. 學生個人寫作（15分），教師巡視，提供個別指導。\n3. 學生小組互讀（5分），圈出覺得最精準的動詞。\n\n【二、分享與回饋】（10分）\n1. 指名 2-3 位學生朗讀作品，全班欣賞。\n2. 教師示範如何給予具體回饋（「你用了___這個動詞，讓我看見___」）。\n3. 學生互評：給同學的作品寫一句欣賞的話。',
                        'time': '30分',
                        'assess': '仿作作品（個人）；口頭回饋觀察',
                    },
                    {
                        'flow': '綜合活動\n（10分）',
                        'activity': '【總結與延伸】\n1. 教師統整三節課的學習歷程：字詞→篇章→修辭→寫作。\n2. 提問：「學完《背影》，你對『表達情感』這件事，有什麼新的理解？」\n3. 布置課後作業：完成 200 字以上的仿作正稿，下次帶來進行同儕評改。\n4. 延伸推薦：余光中《鄉愁》、龍應台《目送》，相同主題的不同情感表達方式。',
                        'time': '10分',
                        'assess': '課後仿作正稿（下節收回）；同儕評改',
                    },
                ],
            },
        ],

        # 素養導向
        'literacy_rows': [
            {
                'dimension': 'A 自主行動',
                'method': '學生自主閱讀課文，標記疑問與感動句子；課後自主完成仿作。',
                'link': '學生透過初讀、精讀過程，建立自主閱讀策略，並將閱讀感受轉化為書寫行動。',
            },
            {
                'dimension': 'B 溝通互動',
                'method': '小組討論動詞分析；仿作朗讀分享；同儕給予口語與書面回饋。',
                'link': '語文表達與傾聽貫穿全課，學生在真實情境中練習溝通協作。',
            },
            {
                'dimension': 'C 社會參與',
                'method': '連結親情與家庭關係；討論現代社會中父子/親子溝通的困境與可能。',
                'link': '課文作為反思「家庭」議題的媒介，引導學生關注自身生命情境中的社會連結。',
            },
        ],

        # 評量
        'assess_rows': [
            {
                'objective': '能正確說出 10 個字詞的注音與字義',
                'method': '字詞學習單（填空、選詞填空）',
                'timing': '第一節課中',
                'criteria': '8 題（含）以上正確為通過；字音、字義各佔 50%',
            },
            {
                'objective': '能分析課文篇章結構',
                'method': '結構分析學習單',
                'timing': '第二節課中',
                'criteria': '能正確標出四次「背影」的位置並說明其作用（可說出 3 次以上）',
            },
            {
                'objective': '能仿照動作描寫方式完成寫作',
                'method': '仿作作品',
                'timing': '第三節課後',
                'criteria': '字數達 100 字；使用至少 3 個精準動詞；不直接點明情感（教師依評量規準評分）',
            },
        ],

        # 延伸
        'extension_rows': [
            {
                'theme': '現代父子情感',
                'resource': '龍應台《目送》（天下文化）',
                'cross': '生命教育、社會',
                'issue': '家庭教育（家）',
            },
            {
                'theme': '鄉愁與記憶',
                'resource': '余光中《鄉愁》（詩）',
                'cross': '社會、歷史',
                'issue': '多元文化（多）',
            },
            {
                'theme': '細節描寫技法',
                'resource': '張曉風散文選（爾雅出版）',
                'cross': '藝術、美感',
                'issue': '美感教育',
            },
        ],
    }

def generate_lesson_plan(args):
    """主要生成函式"""
    data = build_sample_data(args)

    doc = Document()

    # ── 頁面設定（A4）──
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── 全文預設字型 ──
    style = doc.styles['Normal']
    style.font.name = '標楷體'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    # ── 頁首 ──
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f'108課綱素養導向教案｜{data["subject"]}｜{data["title"]}｜{data["grade"]}'
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.runs[0].font.size = Pt(9)
    header_para.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # ── 頁尾（頁碼）──
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar3)

    # ── 封面 ──
    add_cover_page(doc, data['subject'], data['title'], data['grade'],
                   data['publisher'], data['periods'],
                   data['teacher'], data.get('school', ''))

    # ── 九大表格 ──
    add_table1_basic(doc, data)
    add_table2_curriculum(doc, data['competencies'],
                          data['performance'], data['content'])
    add_table3_objectives(doc, data['obj_cog'], data['obj_aff'], data['obj_ski'])
    add_table4_analysis(doc, data['text_type'], data['theme'],
                        data['structure'], data['technique'], data['culture'])
    add_table5_vocab(doc, data['vocab'])
    add_table6_activities(doc, data['periods_data'])
    add_table7_literacy(doc, data['literacy_rows'])
    add_table8_assessment(doc, data['assess_rows'])
    add_table9_extension(doc, data['extension_rows'])

    # ── 儲存 ──
    output = args.output or f'{data["title"]}_教案.docx'
    doc.save(output)
    print(f'✓ 教案已儲存：{output}')
    return output

def main():
    parser = argparse.ArgumentParser(description='臺灣 108 課綱教案生成器')
    parser.add_argument('--subject',   default='國語文', help='科目')
    parser.add_argument('--title',     default='背影',   help='課文/主題名稱')
    parser.add_argument('--grade',     default='國中八年級', help='年級')
    parser.add_argument('--publisher', default='翰林',   help='出版社版本')
    parser.add_argument('--periods',   type=int, default=3, help='教學節數')
    parser.add_argument('--teacher',   default='',      help='設計者姓名')
    parser.add_argument('--school',    default='',      help='學校名稱')
    parser.add_argument('--output',    default='',      help='輸出路徑')
    args = parser.parse_args()
    generate_lesson_plan(args)

if __name__ == '__main__':
    main()
